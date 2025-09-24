from autohive_integrations_sdk import (
    Integration, ExecutionContext, ActionHandler
)
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
import uuid
import os
import base64
from io import BytesIO
from PIL import Image
import markdown
from bs4 import BeautifulSoup
import re

doc_maker = Integration.load()

documents = {}
uploaded_images = {}

def process_files(files: List[Dict[str, Any]]) -> Dict[str, BytesIO]:
    """Process files from the files parameter and return streams by filename"""
    processed_files = {}
    if files:
        for file_item in files:
            content_as_string = file_item['content']

            padded_content_string = content_as_string + '=' * (-len(content_as_string) % 4)

            file_binary_data = base64.urlsafe_b64decode(padded_content_string.encode('ascii'))
            file_stream = BytesIO(file_binary_data)

            processed_files[file_item['name']] = file_stream

    return processed_files

def load_document_from_files(document_id: str, files: List[Dict[str, Any]]) -> None:
    """Load document from files parameter if not in memory"""
    if document_id not in documents and files:
        processed_files = process_files(files)

        for filename, file_stream in processed_files.items():
            if filename.lower().endswith('.docx') or filename.lower().endswith('.bin'):
                try:
                    doc = Document(file_stream)
                    documents[document_id] = doc
                    return
                except Exception as e:
                    continue

        # If no valid Word file found, provide better error message
        available_files = list(processed_files.keys())
        raise ValueError(f"No valid Word file found in files. Tried to load: {available_files}. Files may be corrupted or not Word format.")
    elif document_id not in documents:
        raise ValueError(f"Document {document_id} not found and no files provided for loading")

async def save_and_return_document(original_result: Dict[str, Any], document_id: str, context: ExecutionContext, custom_filename: str = None) -> Dict[str, Any]:
    """Helper to save document and return combined result"""
    save_action = SaveDocumentAction()

    if custom_filename:
        if not custom_filename.lower().endswith('.docx'):
            custom_filename += '.docx'
        file_path = custom_filename
    else:
        file_path = f"{document_id}.docx"

    save_inputs = {
        "document_id": document_id,
        "file_path": file_path
    }
    save_result = await save_action.execute(save_inputs, context)

    combined_result = original_result.copy()
    combined_result.update({
        "saved": save_result["saved"],
        "file_path": save_result["file_path"],
        "file": save_result["file"],
        "error": save_result.get("error", "")
    })
    return combined_result

def hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def parse_markdown_to_docx(doc: Document, markdown_text: str) -> None:
    """Parse markdown text and add elements to Word document"""
    # Convert markdown to HTML
    html = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')

    # Process each HTML element in order
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'table', 'pre']):
        if element.name.startswith('h'):
            # Handle headings
            level = int(element.name[1])  # Extract number from h1, h2, etc.
            text = element.get_text().strip()
            if text:
                doc.add_heading(text, level=level)

        elif element.name == 'p':
            # Handle paragraphs with inline formatting
            paragraph = doc.add_paragraph()
            _add_formatted_text_to_paragraph(paragraph, element)

        elif element.name in ['ul', 'ol']:
            # Handle lists
            is_numbered = element.name == 'ol'
            for li in element.find_all('li', recursive=False):
                text = li.get_text().strip()
                if text:
                    if is_numbered:
                        paragraph = doc.add_paragraph(text, style='List Number')
                    else:
                        paragraph = doc.add_paragraph(text, style='List Bullet')

        elif element.name == 'blockquote':
            # Handle blockquotes
            text = element.get_text().strip()
            if text:
                paragraph = doc.add_paragraph(text, style='Quote')

        elif element.name == 'table':
            # Handle tables
            _add_table_from_html(doc, element)

        elif element.name == 'pre':
            # Handle code blocks
            code_text = element.get_text()
            if code_text:
                paragraph = doc.add_paragraph(code_text)
                # Apply monospace font to code
                for run in paragraph.runs:
                    run.font.name = 'Courier New'

def _add_formatted_text_to_paragraph(paragraph, html_element):
    """Add formatted text from HTML element to Word paragraph"""
    # Handle direct text and formatting
    for content in html_element.contents:
        if hasattr(content, 'name') and content.name:
            # This is an HTML tag
            if content.name == 'strong' or content.name == 'b':
                run = paragraph.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em' or content.name == 'i':
                run = paragraph.add_run(content.get_text())
                run.italic = True
            elif content.name == 'code':
                run = paragraph.add_run(content.get_text())
                run.font.name = 'Courier New'
            elif content.name == 'u':
                run = paragraph.add_run(content.get_text())
                run.underline = True
            else:
                # Nested elements - recursively process only if it has contents
                if hasattr(content, 'contents'):
                    _add_formatted_text_to_paragraph(paragraph, content)
                else:
                    # Just add the text content
                    text = content.get_text()
                    if text:
                        paragraph.add_run(text)
        else:
            # This is plain text (NavigableString) - preserve ALL whitespace
            text = str(content)
            if text:  # Don't strip() here to preserve spaces
                paragraph.add_run(text)

def _add_table_from_html(doc: Document, table_element):
    """Add a table from HTML table element to Word document"""
    rows = table_element.find_all('tr')
    if not rows:
        return

    # Determine table dimensions
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols)

    if max_cols == 0:
        return

    # Create Word table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    # Fill table data
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            if col_idx < max_cols:
                word_cell = table.cell(row_idx, col_idx)
                text = cell.get_text().strip()
                word_cell.text = text

                # Make header cells bold
                if cell.name == 'th':
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

# ---- Action Handlers ----

@doc_maker.action("create_document")
class CreateDocumentAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        title = inputs.get("title")
        markdown_content = inputs.get("markdown_content")
        files = inputs.get("files", [])
        custom_filename = inputs.get("custom_filename")

        processed_files = process_files(files)

        template_file = None
        for filename, file_stream in processed_files.items():
            if filename.lower().endswith('.docx'):
                template_file = file_stream
                break

        if template_file:
            doc = Document(template_file)
        else:
            doc = Document()

        # Add title if provided (but not if we have markdown content with its own title)
        if title and not markdown_content:
            title_paragraph = doc.add_heading(title, level=1)

        # Process markdown content if provided
        if markdown_content:
            parse_markdown_to_docx(doc, markdown_content)

        # Generate unique ID and store document
        document_id = str(uuid.uuid4())
        documents[document_id] = doc

        result = {
            "document_id": document_id,
            "paragraph_count": len(doc.paragraphs),
            "markdown_processed": bool(markdown_content)
        }

        return await save_and_return_document(result, document_id, context, custom_filename)



@doc_maker.action("add_table")
class AddTableAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        rows = inputs["rows"]
        cols = inputs["cols"]
        data = inputs.get("data", [])
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        # Fill table with data if provided
        for row_idx, row_data in enumerate(data[:rows]):
            for col_idx, cell_value in enumerate(row_data[:cols]):
                table.cell(row_idx, col_idx).text = str(cell_value)

        original_result = {"table_rows": rows, "table_cols": cols}
        return await save_and_return_document(original_result, document_id, context)

@doc_maker.action("add_image")
class AddImageAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        width = inputs.get("width")  # in inches
        height = inputs.get("height")  # in inches
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        processed_files = process_files(files)
        image_file = None

        for file_item in files:
            filename = file_item['name']
            content_type = file_item.get('contentType', '')

            # Check if it's an image by extension or content type
            is_image_by_extension = any(filename.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'])
            is_image_by_content_type = content_type.startswith('image/')

            if is_image_by_extension or is_image_by_content_type:
                image_file = processed_files[filename]
                break

        if not image_file:
            raise ValueError("No image file found in files parameter")

        doc = documents[document_id]
        paragraph = doc.add_paragraph()

        if width and height:
            paragraph.add_run().add_picture(image_file, width=Inches(width), height=Inches(height))
        elif width:
            paragraph.add_run().add_picture(image_file, width=Inches(width))
        elif height:
            paragraph.add_run().add_picture(image_file, height=Inches(height))
        else:
            paragraph.add_run().add_picture(image_file)

        original_result = {"image_added": True}
        return await save_and_return_document(original_result, document_id, context)

@doc_maker.action("add_markdown_content")
class AddMarkdownContentAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        markdown_text = inputs["markdown_content"]
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]

        # Parse and add markdown content to document
        parse_markdown_to_docx(doc, markdown_text)

        # Count elements added (approximate by counting paragraphs and headings added)
        elements_added = len([p for p in doc.paragraphs if p.text.strip()])

        original_result = {
            "markdown_processed": True,
            "elements_added": elements_added
        }
        return await save_and_return_document(original_result, document_id, context)

@doc_maker.action("save_document")
class SaveDocumentAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        file_path = inputs["file_path"]

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]

        # Save document to memory buffer instead of disk
        try:
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            file_content = buffer.getvalue()

            # Encode as base64
            content_base64 = base64.b64encode(file_content).decode('utf-8')

            # Get file name from path
            file_name = os.path.basename(file_path)

            return {
                "saved": True,
                "file_path": file_path,
                "file": {
                    "content": content_base64,
                    "name": file_name,
                    "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                }
            }
        except Exception as e:
            return {
                "saved": False,
                "file_path": file_path,
                "file": {
                    "content": "",
                    "name": os.path.basename(file_path),
                    "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                },
                "error": f"Could not generate document for streaming: {str(e)}"
            }

@doc_maker.action("add_page_break")
class AddPageBreakAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)

        original_result = {"page_break_added": True}
        return await save_and_return_document(original_result, document_id, context)

