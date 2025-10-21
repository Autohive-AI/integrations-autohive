from autohive_integrations_sdk import (
    Integration, ExecutionContext, ActionHandler
)
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import uuid
import os
import base64
from io import BytesIO
import markdown
from bs4 import BeautifulSoup
import re
import json

doc_maker = Integration.load()

documents = {}


def process_files(files: List[Dict[str, Any]]) -> Dict[str, BytesIO]:
    """Process files from the files parameter and return streams by filename"""
    processed_files = {}
    if files:
        for file_item in files:
            content_as_string = file_item['content']

            padded_content_string = content_as_string + '=' * (-len(content_as_string) % 4)

            file_binary_data = base64.urlsafe_b64decode(padded_content_string.encode('ascii'))
            file_stream = BytesIO(file_binary_data)

            processed_files[file_item['name']] = file_stream

    return processed_files

def load_document_from_files(document_id: str, files: List[Dict[str, Any]]) -> None:
    """Load document from files parameter if not in memory"""
    if document_id not in documents and files:
        processed_files = process_files(files)

        for filename, file_stream in processed_files.items():
            if filename.lower().endswith('.docx') or filename.lower().endswith('.bin'):
                try:
                    doc = Document(file_stream)
                    documents[document_id] = doc
                    return
                except Exception as e:
                    continue

        # If no valid Word file found, provide better error message
        available_files = list(processed_files.keys())
        raise ValueError(f"No valid Word file found in files. Tried to load: {available_files}. Files may be corrupted or not Word format.")
    elif document_id not in documents:
        raise ValueError(f"Document {document_id} not found and no files provided for loading")

async def save_and_return_document(original_result: Dict[str, Any], document_id: str, context: ExecutionContext, custom_filename: str = None) -> Dict[str, Any]:
    """Helper to save document and return combined result"""
    save_action = SaveDocumentAction()

    if custom_filename:
        if not custom_filename.lower().endswith('.docx'):
            custom_filename += '.docx'
        file_path = custom_filename
    else:
        file_path = f"{document_id}.docx"

    save_inputs = {
        "document_id": document_id,
        "file_path": file_path
    }
    save_result = await save_action.execute(save_inputs, context)

    combined_result = original_result.copy()
    combined_result.update({
        "saved": save_result["saved"],
        "file_path": save_result["file_path"],
        "file": save_result["file"],
        "error": save_result.get("error", "")
    })
    return combined_result

def iter_block_items(parent):
    """
    Iterate through paragraphs and tables in document order.
    Generates a reference to each paragraph and table child within parent.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Parent must be Document or _Cell")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def detect_placeholder_patterns(text: str) -> tuple[bool, str]:
    """Enhanced detection with pattern classification for better LLM optimization"""
    if not text or len(text.strip()) == 0:
        return True, "empty"

    original_text = text.strip()
    text_lower = original_text.lower()

    # Enhanced placeholder pattern detection with classification
    pattern_categories = {
        "formal_placeholder": [
            r'\{\{.*?\}\}',      # {{FIELD}}
            r'\{.*?\}',          # {FIELD}
            r'\[.*?\]',          # [FIELD]
            r'__.*?__',          # __FIELD__
        ],
        "instruction_text": [
            r'\(note:.*?\)',           # (Note: instruction text)
            r'\(delete.*?\)',          # (Delete this section)
            r'\(add.*?\)',             # (Add details here)
            r'\(provide.*?\)',         # (Provide information)
            r'\(insert.*?\)',          # (Insert content here)
            r'\(complete.*?\)',        # (Complete this section)
            r'please (add|insert|provide|enter|complete)', # Please add details
            r'(add|insert|provide|enter|type).*here',      # Insert details here
        ],
        "form_style": [
            r'\w+:\s*[_\-\.]{2,}',     # "Name: ____", "Date: ---"
            r'\w+:\s*\$[_\-\.]+',      # "Amount: $____"
            r'\w+:\s*\[.*?\]',         # "Title: [placeholder]"
            r'\w+:\s*\{.*?\}',         # "Field: {value}"
            r'.*:\s*(tbd|tbc|xxx)',    # "Status: TBD"
        ],
        "business_placeholder": [
            r'(company|client|customer)\s+(name|details|info)', # "company name", "client details"
            r'(project|report|document)\s+(title|name)',        # "project title"
            r'(start|end|due)\s+(date|time)',                   # "start date", "due date"
            r'(total|sub|grand)\s+(amount|cost|price)',         # "total amount"
            r'(contact|manager|author)\s+(name|details)',       # "contact name"
            r'(address|location|venue)\s+(details|info)',       # "address details"
        ],
        "generic_placeholder": [
            r'^(xxx+|yyy+|zzz+|aaa+)$',    # "XXX", "YYY", etc.
            r'^(_+|\.+|-+|\*+)$',          # "___", "...", "---", "***"
            r'^(tbd|tbc|pending|todo)$',   # "TBD", "TBC", "pending"
            r'^(sample|example|dummy|test|placeholder).*', # sample text
            r'lorem ipsum',                 # Lorem ipsum text
            r'(enter|type|add|insert)\s+(text|content|details|info)', # "enter text"
        ],
        "natural_language": [
            r'(data|information|details|content|text)\s+here',     # "data here"
            r'(your|the|this)\s+(name|title|date|info|details)',   # "your name"
            r'(fill|complete|update)\s+(this|here)',               # "fill this"
            r'goes here',                                           # "content goes here"
            r'to be (added|inserted|completed|provided)',          # "to be added"
            r'will be (added|provided|completed)',                 # "will be provided"
        ],
        "short_generic": []  # Handled separately
    }

    # Check each pattern category
    for pattern_type, patterns in pattern_categories.items():
        for pattern in patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True, pattern_type

    # Enhanced short text detection with context
    if len(original_text) < 30:
        # Single words that are likely placeholders in business contexts
        single_word_placeholders = [
            'name', 'title', 'date', 'time', 'amount', 'cost', 'price', 'total',
            'address', 'phone', 'email', 'company', 'client', 'project',
            'description', 'details', 'notes', 'status', 'type', 'category'
        ]

        if text_lower in single_word_placeholders:
            return True, "single_word"

        # Short phrases with business keywords
        business_keywords = ['name', 'date', 'title', 'content', 'text', 'data', 'info', 'details', 'amount', 'cost', 'description']
        if len(original_text) < 20 and any(word in text_lower for word in business_keywords):
            return True, "short_business"

    # Numbers and currency that look like placeholders
    if re.match(r'^[\$£€]?[0-9,\.\s\$]+$', original_text.strip()):
        if len(original_text) < 15:  # Short monetary or numeric values
            return True, "numeric_placeholder"

    # Default patterns that look template-like
    template_indicators = [
        r'(example|sample|demo)',
        r'[a-zA-Z]\s*[a-zA-Z]\s*[a-zA-Z]$',  # "A B C" pattern
        r'^[A-Z\s]{2,10}$',                   # "TITLE TEXT" all caps short text
    ]

    for pattern in template_indicators:
        if re.search(pattern, text_lower):
            return True, "template_indicator"

    return False, "content"

def parse_and_apply_markdown_formatting(target, text: str):
    """
    Centralized markdown parsing and application for any target (paragraph or cell).
    Supports: **bold**, *italic*, `code`, ~~strikethrough~~, __underline__, ***bold italic***, \n line breaks
    """
    # Determine target type and get paragraph to work with
    if hasattr(target, 'clear'):  # This is a paragraph
        paragraph = target
        paragraph.clear()
    elif hasattr(target, 'paragraphs'):  # This is a table cell
        target.text = ""  # Clear cell
        paragraph = target.paragraphs[0] if target.paragraphs else target.add_paragraph()
    else:
        raise ValueError("Target must be a paragraph or table cell")

    # Handle line breaks first by splitting into separate paragraphs
    lines = text.split('\n')

    for line_idx, line in enumerate(lines):
        if line_idx > 0:
            # Add line break for subsequent lines
            paragraph.add_run().add_break()

        if not line.strip():
            continue  # Skip empty lines

        # Enhanced markdown pattern matching (order matters for nested formatting)
        # Match from most specific to least specific
        formatting_patterns = [
            (r'\*\*\*(.*?)\*\*\*', {'bold': True, 'italic': True}),     # ***bold italic***
            (r'\*\*(.*?)\*\*', {'bold': True}),                         # **bold**
            (r'\*(.*?)\*', {'italic': True}),                           # *italic*
            (r'~~(.*?)~~', {'strike': True}),                           # ~~strikethrough~~
            (r'__(.*?)__', {'underline': True}),                        # __underline__
            (r'`(.*?)`', {'code': True}),                               # `code`
        ]

        # Process the line with all formatting patterns
        remaining_text = line
        processed_parts = []

        while remaining_text:
            earliest_match = None
            earliest_pos = len(remaining_text)
            earliest_pattern = None

            # Find the earliest formatting pattern in the remaining text
            for pattern, formatting in formatting_patterns:
                match = re.search(pattern, remaining_text)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start()
                    earliest_pattern = formatting

            if earliest_match:
                # Add text before the match as normal text
                if earliest_pos > 0:
                    processed_parts.append({
                        'text': remaining_text[:earliest_pos],
                        'formatting': {}
                    })

                # Add the formatted text
                processed_parts.append({
                    'text': earliest_match.group(1),
                    'formatting': earliest_pattern
                })

                # Continue with text after the match
                remaining_text = remaining_text[earliest_match.end():]
            else:
                # No more formatting, add remaining text as normal
                if remaining_text:
                    processed_parts.append({
                        'text': remaining_text,
                        'formatting': {}
                    })
                break

        # Apply all processed parts to the paragraph
        for part in processed_parts:
            if part['text']:  # Only add non-empty text
                run = paragraph.add_run(part['text'])

                # Apply formatting
                formatting = part['formatting']
                if formatting.get('bold'):
                    run.bold = True
                if formatting.get('italic'):
                    run.italic = True
                if formatting.get('strike'):
                    run.font.strike = True
                if formatting.get('underline'):
                    run.underline = True
                if formatting.get('code'):
                    run.font.name = 'Courier New'

def has_markdown_formatting(text: str) -> bool:
    """Check if text contains any markdown formatting markers"""
    markers = ['**', '*', '`', '~~', '__', '\n']
    return any(marker in text for marker in markers)

def is_likely_placeholder_context(text: str, find_word: str) -> bool:
    """Determine if a match appears in a placeholder context vs content text"""
    text = text.strip().lower()
    find_word = find_word.lower()

    # Standalone word (likely placeholder)
    if text == find_word:
        return True

    # Form field pattern (word followed by colon and placeholder indicators)
    if re.match(rf'^{re.escape(find_word)}:\s*([_\-\.\[\{{].*)?$', text):
        return True

    # Surrounded by placeholder indicators
    placeholder_indicators = ['{', '}', '[', ']', '_', '-', '.', '(', ')']
    text_around = text.replace(find_word, '').strip()
    if len(text_around) < 10 and any(indicator in text_around for indicator in placeholder_indicators):
        return True

    # In obvious placeholder phrases
    placeholder_phrases = ['insert', 'add', 'enter', 'type', 'provide', 'placeholder', 'here', 'tbd', 'tbc']
    if any(phrase in text for phrase in placeholder_phrases):
        return True

    # Otherwise, likely content text
    return False

def analyze_replacement_safety(find_text: str, matches_found: list) -> dict:
    """Analyze replacement safety and provide intelligent guidance"""
    safe_matches = []
    unsafe_matches = []

    for match in matches_found:
        if is_likely_placeholder_context(match["content"], find_text):
            safe_matches.append(match)
        else:
            unsafe_matches.append(match)

    # Generate intelligent guidance based on actual content
    guidance = []
    alternatives = []

    if len(safe_matches) > 0 and len(unsafe_matches) > 0:
        guidance.append(f"Found {len(safe_matches)} safe placeholders and {len(unsafe_matches)} content text matches")

        # Suggest safer alternatives based on actual safe matches
        safe_contexts = []
        for match in safe_matches[:3]:  # Look at first 3 safe matches
            context = match["content"]
            if ":" in context:
                # Extract the pattern around the colon
                safer_phrase = context.split(find_text)[0] + find_text + ":"
                if safer_phrase not in safe_contexts:
                    safe_contexts.append(safer_phrase.strip())

        if safe_contexts:
            alternatives.extend([f"Use '{ctx}' to target form fields" for ctx in safe_contexts[:2]])

    elif len(unsafe_matches) > 0:
        guidance.append(f"All {len(unsafe_matches)} matches appear to be in content text - very risky")
        alternatives.append(f"Use position updates instead of text replacement")

    elif len(safe_matches) > 0:
        guidance.append(f"All {len(safe_matches)} matches appear to be placeholders - relatively safe")
        if len(safe_matches) > 1:
            alternatives.append(f"Add replace_all=true to confirm you want all {len(safe_matches)} instances replaced")

    return {
        "safety_level": "high_risk" if len(unsafe_matches) > len(safe_matches) else "moderate_risk" if unsafe_matches else "low_risk",
        "safe_matches": len(safe_matches),
        "unsafe_matches": len(unsafe_matches),
        "guidance": guidance,
        "alternatives": alternatives,
        "match_details": [
            {
                "location": f"P{match['index']}" if match["type"] == "paragraph" else f"T{match['table_index']}R{match['row']}C{match['col']}",
                "context": match["content"][:50] + "..." if len(match["content"]) > 50 else match["content"],
                "safety": "SAFE" if match in safe_matches else "RISKY"
            }
            for match in matches_found[:5]  # Show first 5 matches
        ]
    }

def analyze_document_structure(doc: Document) -> dict:
    """Analyze document structure and identify fillable elements"""
    elements = []
    element_index = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            is_fillable, pattern_type = detect_placeholder_patterns(text)

            elements.append({
                "type": "paragraph",
                "index": element_index,
                "content": text,
                "is_fillable": bool(is_fillable),  # Ensure boolean type
                "pattern_type": pattern_type,
                "length": len(text),
                "style": block.style.name if block.style else "Normal"
            })

        elif isinstance(block, Table):
            table_info = {
                "type": "table",
                "index": element_index,
                "rows": len(block.rows),
                "cols": len(block.columns) if block.rows else 0,
                "cells": []
            }

            for row_idx, row in enumerate(block.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    is_fillable, pattern_type = detect_placeholder_patterns(cell_text)

                    table_info["cells"].append({
                        "row": row_idx,
                        "col": col_idx,
                        "content": cell_text,
                        "is_fillable": bool(is_fillable),  # Ensure boolean type
                        "pattern_type": pattern_type,
                        "length": len(cell_text)
                    })

            elements.append(table_info)

        element_index += 1

    # Summary statistics
    fillable_paragraphs = len([e for e in elements if e["type"] == "paragraph" and e["is_fillable"]])
    fillable_cells = sum(len([c for c in e.get("cells", []) if c["is_fillable"]]) for e in elements if e["type"] == "table")

    return {
        "total_elements": len(elements),
        "paragraphs": len([e for e in elements if e["type"] == "paragraph"]),
        "tables": len([e for e in elements if e["type"] == "table"]),
        "fillable_paragraphs": fillable_paragraphs,
        "fillable_cells": fillable_cells,
        "elements": elements
    }

def parse_markdown_to_docx(doc: Document, markdown_text: str) -> None:
    """Parse markdown text and add elements to Word document"""
    # Convert markdown to HTML
    html = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')

    # Process each HTML element in order
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'table', 'pre']):
        if element.name.startswith('h'):
            # Handle headings
            level = int(element.name[1])  # Extract number from h1, h2, etc.
            text = element.get_text().strip()
            if text:
                doc.add_heading(text, level=level)

        elif element.name == 'p':
            # Handle paragraphs with inline formatting
            paragraph = doc.add_paragraph()
            _add_formatted_text_to_paragraph(paragraph, element)

        elif element.name in ['ul', 'ol']:
            # Handle lists
            is_numbered = element.name == 'ol'
            for li in element.find_all('li', recursive=False):
                text = li.get_text().strip()
                if text:
                    if is_numbered:
                        paragraph = doc.add_paragraph(text, style='List Number')
                    else:
                        paragraph = doc.add_paragraph(text, style='List Bullet')

        elif element.name == 'blockquote':
            # Handle blockquotes
            text = element.get_text().strip()
            if text:
                paragraph = doc.add_paragraph(text, style='Quote')

        elif element.name == 'table':
            # Handle tables
            _add_table_from_html(doc, element)

        elif element.name == 'pre':
            # Handle code blocks
            code_text = element.get_text()
            if code_text:
                paragraph = doc.add_paragraph(code_text)
                # Apply monospace font to code
                for run in paragraph.runs:
                    run.font.name = 'Courier New'

def _add_formatted_text_to_paragraph(paragraph, html_element):
    """Add formatted text from HTML element to Word paragraph"""
    # Handle direct text and formatting
    for content in html_element.contents:
        if hasattr(content, 'name') and content.name:
            # This is an HTML tag
            if content.name == 'strong' or content.name == 'b':
                run = paragraph.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em' or content.name == 'i':
                run = paragraph.add_run(content.get_text())
                run.italic = True
            elif content.name == 'code':
                run = paragraph.add_run(content.get_text())
                run.font.name = 'Courier New'
            elif content.name == 'u':
                run = paragraph.add_run(content.get_text())
                run.underline = True
            else:
                # Nested elements - recursively process only if it has contents
                if hasattr(content, 'contents'):
                    _add_formatted_text_to_paragraph(paragraph, content)
                else:
                    # Just add the text content
                    text = content.get_text()
                    if text:
                        paragraph.add_run(text)
        else:
            # This is plain text (NavigableString) - preserve ALL whitespace
            text = str(content)
            if text:  # Don't strip() here to preserve spaces
                paragraph.add_run(text)

def _add_table_from_html(doc: Document, table_element):
    """Add a table from HTML table element to Word document"""
    rows = table_element.find_all('tr')
    if not rows:
        return

    # Determine table dimensions
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols)

    if max_cols == 0:
        return

    # Create Word table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    # Fill table data
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            if col_idx < max_cols:
                word_cell = table.cell(row_idx, col_idx)
                text = cell.get_text().strip()
                word_cell.text = text

                # Make header cells bold
                if cell.name == 'th':
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

# ---- Action Handlers ----

@doc_maker.action("get_document_elements")
class GetDocumentElementsAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]

        # Analyze document structure
        analysis = analyze_document_structure(doc)

        # Always return LLM-optimized response (fillable elements only)
        fillable_paragraphs = []
        fillable_cells = []
        pattern_counts = {}

        for element in analysis["elements"]:
            if element["type"] == "paragraph" and bool(element["is_fillable"]):
                fillable_paragraphs.append({
                    "id": f"p{element['index']}",
                    "content": str(element["content"]),
                    "pattern": str(element["pattern_type"]),
                    "style": str(element["style"])
                })
                # Count patterns
                pattern = element["pattern_type"]
                pattern_counts[pattern] = pattern_counts.get(pattern, 0) + 1

            elif element["type"] == "table":
                for cell in element["cells"]:
                    if bool(cell["is_fillable"]):
                        fillable_cells.append({
                            "id": f"t{element['index']}r{cell['row']}c{cell['col']}",
                            "content": str(cell["content"]),
                            "pattern": str(cell["pattern_type"]),
                            "location": f"Table {element['index']}, Row {cell['row']}, Col {cell['col']}"
                        })
                        # Count patterns
                        pattern = cell["pattern_type"]
                        pattern_counts[pattern] = pattern_counts.get(pattern, 0) + 1

        return {
            "template_summary": {
                "structure": f"{analysis['paragraphs']}p,{analysis['tables']}t",
                "fillable_total": int(analysis["fillable_paragraphs"] + analysis["fillable_cells"]),
                "content_elements_hidden": int(analysis["total_elements"] - len(fillable_paragraphs) - len(fillable_cells))
            },
            "fillable_paragraphs": fillable_paragraphs,
            "fillable_cells": fillable_cells,
            "pattern_distribution": pattern_counts,
            "recommended_strategy": "mixed" if len(pattern_counts) > 2 else "single_method",
            "template_ready": True
        }

@doc_maker.action("create_document")
class CreateDocumentAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        title = inputs.get("title")
        markdown_content = inputs.get("markdown_content")
        files = inputs.get("files", [])
        custom_filename = inputs.get("custom_filename")

        processed_files = process_files(files)

        template_file = None
        for filename, file_stream in processed_files.items():
            if filename.lower().endswith('.docx'):
                template_file = file_stream
                break

        if template_file:
            doc = Document(template_file)
        else:
            doc = Document()

        # Add title if provided (but not if we have markdown content with its own title)
        if title and not markdown_content:
            title_paragraph = doc.add_heading(title, level=1)

        # Process markdown content if provided
        if markdown_content:
            parse_markdown_to_docx(doc, markdown_content)

        # Generate unique ID and store document
        document_id = str(uuid.uuid4())
        documents[document_id] = doc

        result = {
            "document_id": document_id,
            "paragraph_count": len(doc.paragraphs),
            "markdown_processed": bool(markdown_content)
        }

        return await save_and_return_document(result, document_id, context, custom_filename)



@doc_maker.action("add_table")
class AddTableAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        rows = inputs["rows"]
        cols = inputs["cols"]
        data = inputs.get("data", [])
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        # Fill table with data if provided
        for row_idx, row_data in enumerate(data[:rows]):
            for col_idx, cell_value in enumerate(row_data[:cols]):
                table.cell(row_idx, col_idx).text = str(cell_value)

        original_result = {"table_rows": rows, "table_cols": cols}
        return await save_and_return_document(original_result, document_id, context)

@doc_maker.action("add_image")
class AddImageAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        width = inputs.get("width")  # in inches
        height = inputs.get("height")  # in inches
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        processed_files = process_files(files)
        image_file = None

        for file_item in files:
            filename = file_item['name']
            content_type = file_item.get('contentType', '')

            # Check if it's an image by extension or content type
            is_image_by_extension = any(filename.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'])
            is_image_by_content_type = content_type.startswith('image/')

            if is_image_by_extension or is_image_by_content_type:
                image_file = processed_files[filename]
                break

        if not image_file:
            raise ValueError("No image file found in files parameter")

        doc = documents[document_id]
        paragraph = doc.add_paragraph()

        if width and height:
            paragraph.add_run().add_picture(image_file, width=Inches(width), height=Inches(height))
        elif width:
            paragraph.add_run().add_picture(image_file, width=Inches(width))
        elif height:
            paragraph.add_run().add_picture(image_file, height=Inches(height))
        else:
            paragraph.add_run().add_picture(image_file)

        original_result = {"image_added": True}
        return await save_and_return_document(original_result, document_id, context)

@doc_maker.action("add_markdown_content")
class AddMarkdownContentAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        markdown_text = inputs["markdown_content"]
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]

        # Parse and add markdown content to document
        parse_markdown_to_docx(doc, markdown_text)

        # Count elements added (approximate by counting paragraphs and headings added)
        elements_added = len([p for p in doc.paragraphs if p.text.strip()])

        original_result = {
            "markdown_processed": True,
            "elements_added": elements_added
        }
        return await save_and_return_document(original_result, document_id, context)

@doc_maker.action("update_by_position")
class UpdateByPositionAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        updates = inputs["updates"]
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]
        changes_made = []

        # Process updates
        for update in updates:
            update_type = update.get("type")

            if update_type == "paragraph":
                paragraph_index = update["index"]
                new_content = update["content"]

                # Get paragraph by index using iter_block_items
                paragraphs = [block for block in iter_block_items(doc) if isinstance(block, Paragraph)]

                if paragraph_index < len(paragraphs):
                    paragraph = paragraphs[paragraph_index]

                    # Preserve formatting by clearing and adding new text
                    paragraph.clear()
                    paragraph.text = new_content
                    changes_made.append(f"Updated paragraph {paragraph_index}")
                else:
                    changes_made.append(f"Paragraph {paragraph_index} not found")

            elif update_type == "table_cell":
                table_index = update["table_index"]
                row = update["row"]
                col = update["col"]
                new_content = update["content"]

                # Get table by index
                tables = [block for block in iter_block_items(doc) if isinstance(block, Table)]

                if table_index < len(tables):
                    table = tables[table_index]
                    if row < len(table.rows) and col < len(table.columns):
                        cell = table.cell(row, col)
                        cell.text = new_content
                        changes_made.append(f"Updated table {table_index} cell ({row},{col})")
                    else:
                        changes_made.append(f"Cell ({row},{col}) out of range in table {table_index}")
                else:
                    changes_made.append(f"Table {table_index} not found")

        # Create LLM-optimized response
        successful_updates = [change for change in changes_made if "Updated" in change]
        failed_updates = [change for change in changes_made if "not found" in change or "out of range" in change]

        original_result = {
            "success": len(successful_updates) > 0,
            "applied": len(successful_updates),
            "failed": len(failed_updates),
            "summary": f"Updated {len(successful_updates)} elements" + (f", {len(failed_updates)} failed" if failed_updates else ""),
            "failures": failed_updates[:3] if failed_updates else []  # Limit failure details
        }
        return await save_and_return_document(original_result, document_id, context)

@doc_maker.action("find_and_replace")
class FindAndReplaceAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]

        # Handle both array and JSON string formats (platform compatibility)
        replacements = inputs["replacements"]
        if isinstance(replacements, str):
            try:
                replacements = json.loads(replacements)
            except json.JSONDecodeError:
                raise ValueError("Invalid replacements format: must be array or valid JSON string")

        case_sensitive = inputs.get("case_sensitive", False)
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]
        total_replacements = 0
        warnings = []
        skipped_replacements = []

        for replacement in replacements:
            find_text = replacement["find"]
            replace_text = replacement.get("replace", "")  # Default to empty string if not provided
            replace_all = replacement.get("replace_all", False)
            remove_paragraph = replacement.get("remove_paragraph", False)

            # Validation
            if not find_text or len(find_text.strip()) == 0:
                warnings.append(f"Skipped replacement: 'find' text cannot be empty")
                continue

            # Handle space-as-delete (convert single space to empty for deletion)
            if replace_text == " ":
                replace_text = ""

            # First, scan for all matches to check for multiple occurrences
            matches_found = []

            # Scan paragraphs for matches
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if case_sensitive:
                    if find_text in paragraph.text:
                        matches_found.append({
                            "type": "paragraph",
                            "index": para_idx,
                            "content": paragraph.text[:100] + "..." if len(paragraph.text) > 100 else paragraph.text,
                            "context": f"Paragraph {para_idx}"
                        })
                else:
                    if find_text.lower() in paragraph.text.lower():
                        matches_found.append({
                            "type": "paragraph",
                            "index": para_idx,
                            "content": paragraph.text[:100] + "..." if len(paragraph.text) > 100 else paragraph.text,
                            "context": f"Paragraph {para_idx}"
                        })

            # Scan tables for matches
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        if case_sensitive:
                            if find_text in cell.text:
                                matches_found.append({
                                    "type": "table_cell",
                                    "table_index": table_idx,
                                    "row": row_idx,
                                    "col": col_idx,
                                    "content": cell.text,
                                    "context": f"Table {table_idx}, Row {row_idx}, Col {col_idx}"
                                })
                        else:
                            if find_text.lower() in cell.text.lower():
                                matches_found.append({
                                    "type": "table_cell",
                                    "table_index": table_idx,
                                    "row": row_idx,
                                    "col": col_idx,
                                    "content": cell.text,
                                    "context": f"Table {table_idx}, Row {row_idx}, Col {col_idx}"
                                })

            # Enhanced safety check for multiple matches
            if len(matches_found) > 1 and not replace_all:
                safety_analysis = analyze_replacement_safety(find_text, matches_found)

                if safety_analysis["safety_level"] == "high_risk":
                    # Block high-risk replacements with detailed analysis
                    skipped_replacement = {
                        "CRITICAL_WARNING": f"BLOCKED '{find_text}' - {safety_analysis['unsafe_matches']} unsafe content matches",
                        "find_phrase": find_text,
                        "risk_assessment": safety_analysis["safety_level"],
                        "safe_placeholders": safety_analysis["safe_matches"],
                        "content_text_matches": safety_analysis["unsafe_matches"],
                        "intelligent_alternatives": safety_analysis["alternatives"],
                        "match_analysis": safety_analysis["match_details"],
                        "fix_required": "Use context-specific phrases from alternatives or position updates"
                    }
                    skipped_replacements.append(skipped_replacement)
                    warnings.append(f"BLOCKED '{find_text}': High risk - {safety_analysis['unsafe_matches']} content matches detected")
                    continue

                else:
                    # Warn about moderate risk but allow
                    warning_details = {
                        "phrase": find_text,
                        "matches": len(matches_found),
                        "risk_level": safety_analysis["safety_level"],
                        "alternatives": safety_analysis["alternatives"],
                        "context_review": safety_analysis["match_details"][:3]
                    }
                    skipped_replacements.append(warning_details)
                    warnings.append(f"Proceeding with '{find_text}' but review recommended - {len(matches_found)} matches")

            elif len(matches_found) == 0:
                warnings.append(f"No matches found for '{find_text}'")
                continue

            # Proceed with replacement
            replacements_count = 0

            # Perform replacements in paragraphs with spacing control
            paragraphs_to_remove = []
            for paragraph in doc.paragraphs:
                text_matches = False
                if case_sensitive:
                    text_matches = find_text in paragraph.text
                else:
                    text_matches = find_text.lower() in paragraph.text.lower()

                if text_matches:
                    original_text = paragraph.text
                    is_full_paragraph_match = original_text.strip() == find_text.strip()

                    if is_full_paragraph_match and replace_text.strip() == "" and remove_paragraph:
                        # Mark paragraph for removal to eliminate spacing
                        paragraphs_to_remove.append(paragraph)
                        replacements_count += 1
                    else:
                        # Normal text replacement (preserves spacing) with enhanced formatting support
                        if case_sensitive:
                            new_text = original_text.replace(find_text, replace_text)
                        else:
                            new_text = re.sub(re.escape(find_text), replace_text, original_text, flags=re.IGNORECASE)

                        # Use centralized parser for enhanced formatting support
                        if has_markdown_formatting(new_text):
                            parse_and_apply_markdown_formatting(paragraph, new_text)
                        else:
                            paragraph.text = new_text

                        replacements_count += 1

            # Remove marked paragraphs
            for paragraph in paragraphs_to_remove:
                try:
                    p = paragraph._element
                    p.getparent().remove(p)
                except:
                    # Fallback: just clear the text
                    paragraph.clear()

            # Perform replacements in tables with enhanced formatting support
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if case_sensitive:
                            if find_text in cell.text:
                                new_text = cell.text.replace(find_text, replace_text)

                                # Use centralized parser for enhanced formatting support
                                if has_markdown_formatting(new_text):
                                    parse_and_apply_markdown_formatting(cell, new_text)
                                else:
                                    cell.text = new_text
                                replacements_count += 1
                        else:
                            original_text = cell.text
                            new_text = re.sub(re.escape(find_text), replace_text, original_text, flags=re.IGNORECASE)
                            if new_text != original_text:
                                # Use centralized parser for enhanced formatting support
                                if has_markdown_formatting(new_text):
                                    parse_and_apply_markdown_formatting(cell, new_text)
                                else:
                                    cell.text = new_text
                                replacements_count += 1

            total_replacements += replacements_count

        # Create LLM-optimized response with proper field handling
        optimized_blocked = []
        for skipped in skipped_replacements:
            # Handle different skipped replacement formats
            if "CRITICAL_WARNING" in skipped:
                # This is from the new safety analysis format
                optimized_blocked.append({
                    "phrase": skipped.get("find_phrase", "unknown"),
                    "warning": skipped.get("CRITICAL_WARNING", ""),
                    "risk": skipped.get("risk_assessment", "unknown"),
                    "safe_matches": skipped.get("safe_placeholders", 0),
                    "unsafe_matches": skipped.get("content_text_matches", 0),
                    "alternatives": skipped.get("intelligent_alternatives", []),
                    "action_required": skipped.get("fix_required", "")
                })
            elif "WARNING" in skipped:
                # This is a moderate risk warning
                optimized_blocked.append({
                    "phrase": skipped.get("phrase", "unknown"),
                    "warning": skipped.get("WARNING", ""),
                    "matches": skipped.get("matches", 0),
                    "alternatives": skipped.get("alternatives", [])
                })
            else:
                # Fallback for any other format
                optimized_blocked.append({
                    "phrase": str(skipped),
                    "warning": "Format error in safety analysis"
                })

        original_result = {
            "success": total_replacements > 0 and len([s for s in skipped_replacements if "CRITICAL_WARNING" in s]) == 0,
            "replaced": total_replacements,
            "processed": len(replacements),
            "blocked": optimized_blocked,
            "alerts": warnings[:3] if warnings else [],
            "safety_active": True
        }
        return await save_and_return_document(original_result, document_id, context)

@doc_maker.action("fill_template_fields")
class FillTemplateFieldsAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        template_data = inputs["template_data"]
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]
        changes_made = []

        # Process different types of template data

        # 1. Placeholder data ({{field}} format)
        if "placeholder_data" in template_data:
            for placeholder, value in template_data["placeholder_data"].items():
                replacement_count = 0

                # Replace in paragraphs with enhanced formatting support
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        new_text = paragraph.text.replace(placeholder, str(value))
                        # Use centralized parser for all replacements
                        if has_markdown_formatting(new_text):
                            parse_and_apply_markdown_formatting(paragraph, new_text)
                        else:
                            paragraph.text = new_text
                        replacement_count += 1

                # Replace in tables with enhanced formatting support
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if placeholder in cell.text:
                                new_text = cell.text.replace(placeholder, str(value))
                                # Use centralized parser for all replacements
                                if has_markdown_formatting(new_text):
                                    parse_and_apply_markdown_formatting(cell, new_text)
                                else:
                                    cell.text = new_text
                                replacement_count += 1

                if replacement_count > 0:
                    changes_made.append(f"Replaced '{placeholder}' {replacement_count} times")

        # 2. Position-based updates
        if "position_data" in template_data:
            paragraphs = [block for block in iter_block_items(doc) if isinstance(block, Paragraph)]
            tables = [block for block in iter_block_items(doc) if isinstance(block, Table)]

            for position_key, new_content in template_data["position_data"].items():
                if position_key.startswith("paragraph_"):
                    idx = int(position_key.split("_")[1])
                    if idx < len(paragraphs):
                        # Use centralized parser for all content
                        if has_markdown_formatting(str(new_content)):
                            parse_and_apply_markdown_formatting(paragraphs[idx], str(new_content))
                        else:
                            paragraphs[idx].text = str(new_content)
                        changes_made.append(f"Updated paragraph {idx}")

                elif position_key.startswith("table_"):
                    # Format: table_0_row_1_col_2
                    parts = position_key.split("_")
                    table_idx = int(parts[1])
                    row_idx = int(parts[3])
                    col_idx = int(parts[5])

                    if table_idx < len(tables):
                        table = tables[table_idx]
                        if row_idx < len(table.rows) and col_idx < len(table.columns):
                            cell = table.cell(row_idx, col_idx)
                            # Use centralized parser for all content
                            if has_markdown_formatting(str(new_content)):
                                parse_and_apply_markdown_formatting(cell, str(new_content))
                            else:
                                cell.text = str(new_content)
                            changes_made.append(f"Updated table {table_idx} cell ({row_idx},{col_idx})")

        # 3. Search and replace patterns (with safety analysis)
        safety_warnings = []
        if "search_replace" in template_data:
            for item in template_data["search_replace"]:
                find_text = item["find"]
                replace_text = item["replace"]
                replace_all = item.get("replace_all", False)
                remove_paragraph = item.get("remove_paragraph", False)

                # First, scan for all matches to analyze safety
                matches_found = []

                # Scan paragraphs
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    if find_text.lower() in paragraph.text.lower():
                        matches_found.append({
                            "type": "paragraph",
                            "index": para_idx,
                            "content": paragraph.text,
                            "context": f"Paragraph {para_idx}"
                        })

                # Scan tables
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            if find_text.lower() in cell.text.lower():
                                matches_found.append({
                                    "type": "table_cell",
                                    "table_index": table_idx,
                                    "row": row_idx,
                                    "col": col_idx,
                                    "content": cell.text,
                                    "context": f"Table {table_idx}, Row {row_idx}, Col {col_idx}"
                                })

                # Analyze safety if multiple matches
                if len(matches_found) > 1 and not replace_all:
                    safety_analysis = analyze_replacement_safety(find_text, matches_found)

                    if safety_analysis["safety_level"] == "high_risk":
                        # Block high-risk replacements
                        safety_warnings.append({
                            "CRITICAL_WARNING": f"BLOCKED replacement of '{find_text}' - {safety_analysis['unsafe_matches']} unsafe matches detected",
                            "find_phrase": find_text,
                            "risk_level": safety_analysis["safety_level"],
                            "safe_matches": safety_analysis["safe_matches"],
                            "unsafe_matches": safety_analysis["unsafe_matches"],
                            "alternatives": safety_analysis["alternatives"],
                            "match_details": safety_analysis["match_details"],
                            "action_required": "Use more specific context or position-based updates"
                        })
                        continue  # Skip this dangerous replacement

                    else:
                        # Warn but allow moderate risk replacements
                        safety_warnings.append({
                            "WARNING": f"'{find_text}' has {len(matches_found)} matches - review recommended",
                            "alternatives": safety_analysis["alternatives"],
                            "match_preview": safety_analysis["match_details"][:3]
                        })

                # Proceed with replacement
                replacement_count = 0
                paragraphs_to_remove = []

                # Replace in paragraphs with spacing control and formatting support
                for paragraph in doc.paragraphs:
                    if find_text.lower() in paragraph.text.lower():
                        original_text = paragraph.text
                        is_full_match = original_text.strip().lower() == find_text.lower()

                        if is_full_match and replace_text.strip() == "" and remove_paragraph:
                            paragraphs_to_remove.append(paragraph)
                            replacement_count += 1
                        else:
                            # Perform text replacement
                            new_text = re.sub(re.escape(find_text), replace_text, original_text, flags=re.IGNORECASE)

                            # Use centralized parser for enhanced formatting support
                            if has_markdown_formatting(new_text):
                                parse_and_apply_markdown_formatting(paragraph, new_text)
                            else:
                                paragraph.text = new_text

                            if new_text != original_text:
                                replacement_count += 1

                # Remove marked paragraphs
                for paragraph in paragraphs_to_remove:
                    try:
                        p = paragraph._element
                        p.getparent().remove(p)
                    except:
                        paragraph.clear()

                # Replace in tables with enhanced formatting support
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if find_text.lower() in cell.text.lower():
                                original_text = cell.text
                                new_text = re.sub(re.escape(find_text), replace_text, original_text, flags=re.IGNORECASE)

                                # Use centralized parser for enhanced formatting support
                                if has_markdown_formatting(new_text):
                                    parse_and_apply_markdown_formatting(cell, new_text)
                                else:
                                    cell.text = new_text

                                if new_text != original_text:
                                    replacement_count += 1

                if replacement_count > 0:
                    changes_made.append(f"Found and replaced '{find_text}' {replacement_count} times")

        # Create LLM-optimized response with prominent safety warnings
        has_critical_warnings = any("CRITICAL_WARNING" in str(warning) for warning in safety_warnings)
        blocked_operations = len([w for w in safety_warnings if "BLOCKED" in str(w)])

        change_summary = {}
        for change in changes_made:
            if "Replaced" in change:
                change_summary["placeholders"] = change_summary.get("placeholders", 0) + 1
            elif "Found and replaced" in change:
                change_summary["searches"] = change_summary.get("searches", 0) + 1
            elif "Updated" in change:
                change_summary["positions"] = change_summary.get("positions", 0) + 1

        original_result = {
            "SAFETY_STATUS": "CRITICAL_ISSUES_DETECTED" if has_critical_warnings else "OK",
            "success": len(changes_made) > 0 and not has_critical_warnings,
            "completed_operations": len(changes_made),
            "blocked_operations": blocked_operations,
            "safety_warnings": safety_warnings,
            "filled_summary": change_summary,
            "template_status": "partially_complete" if blocked_operations > 0 else "complete",
            "action_required": "Review safety warnings and use more specific context" if safety_warnings else "none"
        }
        return await save_and_return_document(original_result, document_id, context)


@doc_maker.action("save_document")
class SaveDocumentAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        file_path = inputs["file_path"]

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]

        # Save document to memory buffer instead of disk
        try:
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            file_content = buffer.getvalue()

            # Encode as base64
            content_base64 = base64.b64encode(file_content).decode('utf-8')

            # Get file name from path
            file_name = os.path.basename(file_path)

            return {
                "saved": True,
                "file_path": file_path,
                "file": {
                    "content": content_base64,
                    "name": file_name,
                    "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                }
            }
        except Exception as e:
            return {
                "saved": False,
                "file_path": file_path,
                "file": {
                    "content": "",
                    "name": os.path.basename(file_path),
                    "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                },
                "error": f"Could not generate document for streaming: {str(e)}"
            }

@doc_maker.action("add_page_break")
class AddPageBreakAction(ActionHandler):
    async def execute(self, inputs: Dict[str, Any], context: ExecutionContext):
        document_id = inputs["document_id"]
        files = inputs.get("files", [])

        load_document_from_files(document_id, files)

        if document_id not in documents:
            raise ValueError(f"Document {document_id} not found")

        doc = documents[document_id]
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)

        original_result = {"page_break_added": True}
        return await save_and_return_document(original_result, document_id, context)

